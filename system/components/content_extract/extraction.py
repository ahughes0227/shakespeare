"""Content extraction, one operator with config-selected backends.

Every backend degrades explicitly.  A missing OCR binary or an absent optional dependency
returns a reason rather than empty text, so the item is quarantined by the `extract`
obligation instead of silently producing a nameless file.

A corpus is extracted by `extract_many`, which runs items in parallel.  This is the first
threaded work in the runtime and the reason ADR 0006 committed to a free-threaded
interpreter: extraction is per-item, share-nothing and by far the largest serial cost in a
run.  Two properties survive the change — results come back in input order, and the
lxml-backed backends still run one at a time on the calling thread.  See `extract_many`.
"""

from __future__ import annotations

import os
from collections.abc import Sequence
from concurrent.futures import ThreadPoolExecutor, as_completed
from enum import StrEnum
from pathlib import Path
from typing import NamedTuple

from pydantic import Field

from ...contracts import Contract


class Backend(StrEnum):
    PDF_TEXT = "pdf_text"
    PDF_OCR = "pdf_ocr"
    DOCX = "docx"
    XLSX = "xlsx"
    EMAIL = "email"
    IMAGE_OCR = "image_ocr"
    #: Try backends in order until one yields text.  Fallback is configuration, not agent
    #: reasoning, because an agent never sees its own operator output.
    AUTO_CHAIN = "auto_chain"


class Span(Contract):
    """Where a piece of text came from, so a field value can cite its provenance."""

    page: int | None = None
    start: int
    end: int


class Extraction(Contract):
    item_id: str
    backend: str
    text: str = ""
    spans: tuple[Span, ...] = ()
    unavailable_reason: str | None = None
    char_count: int = 0

    @property
    def usable(self) -> bool:
        return bool(self.text.strip())


class ExtractOptions(Contract):
    page_limit: int = Field(default=20, ge=1, le=500)
    char_limit: int = Field(default=200_000, ge=100)


class Item(NamedTuple):
    """One unit of extraction work.

    A tuple rather than a `Contract` because it carries an absolute path, which is a
    property of this machine and belongs in no plan, journal or telemetry envelope.
    """

    item_id: str
    path: Path
    media_type: str = "application/octet-stream"


#: Ordered fallback for `auto_chain`, chosen by media type.
_CHAIN: dict[str, tuple[Backend, ...]] = {
    "application/pdf": (Backend.PDF_TEXT, Backend.PDF_OCR),
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": (Backend.DOCX,),
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": (Backend.XLSX,),
    "message/rfc822": (Backend.EMAIL,),
    "application/vnd.ms-outlook": (Backend.EMAIL,),
    "image/png": (Backend.IMAGE_OCR,),
    "image/jpeg": (Backend.IMAGE_OCR,),
    "image/tiff": (Backend.IMAGE_OCR,),
    "text/plain": (Backend.PDF_TEXT,),
}


def _unavailable(item_id: str, backend: Backend, reason: str) -> Extraction:
    return Extraction(item_id=item_id, backend=str(backend), unavailable_reason=reason)


def _clip(text: str, options: ExtractOptions) -> str:
    return text[: options.char_limit]


def _pdf_text(item_id: str, path: Path, options: ExtractOptions) -> Extraction:
    try:
        import pdfplumber
    except ImportError:
        return _unavailable(item_id, Backend.PDF_TEXT, "backend_unavailable:pdfplumber")

    pages: list[str] = []
    spans: list[Span] = []
    cursor = 0
    try:
        with pdfplumber.open(path) as document:
            for number, page in enumerate(document.pages[: options.page_limit], start=1):
                text = page.extract_text() or ""
                if text:
                    spans.append(Span(page=number, start=cursor, end=cursor + len(text)))
                    cursor += len(text) + 1
                    pages.append(text)
    except Exception as exc:  # noqa: BLE001 - a corrupt document is data, not a crash
        return _unavailable(item_id, Backend.PDF_TEXT, f"unreadable:{type(exc).__name__}")

    text = _clip("\n".join(pages), options)
    if not text.strip():
        # A scanned PDF has no text layer.  Saying so lets auto_chain reach for OCR.
        return _unavailable(item_id, Backend.PDF_TEXT, "no_text_layer")
    return Extraction(
        item_id=item_id,
        backend=str(Backend.PDF_TEXT),
        text=text,
        spans=tuple(spans),
        char_count=len(text),
    )


def _ocr_available() -> str | None:
    try:
        import pytesseract
    except ImportError:
        return "backend_unavailable:pytesseract"
    try:
        pytesseract.get_tesseract_version()
    except Exception:  # noqa: BLE001 - the binary is absent or not on PATH
        return "ocr_unavailable:tesseract_binary_missing"
    return None


def _image_ocr(item_id: str, path: Path, options: ExtractOptions) -> Extraction:
    reason = _ocr_available()
    if reason:
        return _unavailable(item_id, Backend.IMAGE_OCR, reason)
    import pytesseract
    from PIL import Image

    try:
        with Image.open(path) as image:
            text = _clip(pytesseract.image_to_string(image), options)
    except Exception as exc:  # noqa: BLE001
        return _unavailable(item_id, Backend.IMAGE_OCR, f"unreadable:{type(exc).__name__}")
    if not text.strip():
        return _unavailable(item_id, Backend.IMAGE_OCR, "ocr_produced_no_text")
    return Extraction(
        item_id=item_id, backend=str(Backend.IMAGE_OCR), text=text, char_count=len(text)
    )


def _pdf_ocr(item_id: str, path: Path, options: ExtractOptions) -> Extraction:
    reason = _ocr_available()
    if reason:
        return _unavailable(item_id, Backend.PDF_OCR, reason)
    try:
        import pdfplumber
        import pytesseract
    except ImportError:
        return _unavailable(item_id, Backend.PDF_OCR, "backend_unavailable:pdfplumber")

    pages: list[str] = []
    try:
        with pdfplumber.open(path) as document:
            for page in document.pages[: options.page_limit]:
                image = page.to_image(resolution=200).original
                pages.append(pytesseract.image_to_string(image))
    except Exception as exc:  # noqa: BLE001
        return _unavailable(item_id, Backend.PDF_OCR, f"unreadable:{type(exc).__name__}")

    text = _clip("\n".join(pages), options)
    if not text.strip():
        return _unavailable(item_id, Backend.PDF_OCR, "ocr_produced_no_text")
    return Extraction(
        item_id=item_id, backend=str(Backend.PDF_OCR), text=text, char_count=len(text)
    )


def _docx(item_id: str, path: Path, options: ExtractOptions) -> Extraction:
    try:
        import docx
    except ImportError:
        return _unavailable(item_id, Backend.DOCX, "backend_unavailable:python-docx")
    try:
        document = docx.Document(str(path))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
    except Exception as exc:  # noqa: BLE001
        return _unavailable(item_id, Backend.DOCX, f"unreadable:{type(exc).__name__}")
    text = _clip("\n".join(parts), options)
    if not text.strip():
        return _unavailable(item_id, Backend.DOCX, "document_has_no_text")
    return Extraction(item_id=item_id, backend=str(Backend.DOCX), text=text, char_count=len(text))


def _xlsx(item_id: str, path: Path, options: ExtractOptions) -> Extraction:
    try:
        import openpyxl
    except ImportError:
        return _unavailable(item_id, Backend.XLSX, "backend_unavailable:openpyxl")
    try:
        workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
        rows: list[str] = []
        for sheet in workbook.worksheets[: options.page_limit]:
            rows.append(f"# {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                rows.append("\t".join("" if cell is None else str(cell) for cell in row))
        workbook.close()
    except Exception as exc:  # noqa: BLE001
        return _unavailable(item_id, Backend.XLSX, f"unreadable:{type(exc).__name__}")
    text = _clip("\n".join(rows), options)
    if not text.strip():
        return _unavailable(item_id, Backend.XLSX, "workbook_has_no_values")
    return Extraction(item_id=item_id, backend=str(Backend.XLSX), text=text, char_count=len(text))


def _email(item_id: str, path: Path, options: ExtractOptions) -> Extraction:
    if path.suffix.lower() == ".msg":
        try:
            import extract_msg
        except ImportError:
            return _unavailable(item_id, Backend.EMAIL, "backend_unavailable:extract-msg")
        try:
            # Read defensively: extract_msg's stubs type the factory as MSGFile, which
            # declares none of these, and the attribute set has moved between releases.
            with extract_msg.Message(str(path)) as outlook:  # type: ignore[no-untyped-call]
                parts = [
                    f"From: {getattr(outlook, 'sender', '')}",
                    f"To: {getattr(outlook, 'to', '')}",
                    f"Subject: {getattr(outlook, 'subject', '')}",
                    f"Date: {getattr(outlook, 'date', '')}",
                    str(getattr(outlook, "body", "") or ""),
                ]
        except Exception as exc:  # noqa: BLE001
            return _unavailable(item_id, Backend.EMAIL, f"unreadable:{type(exc).__name__}")
    else:
        import email
        from email import policy

        try:
            parsed = email.message_from_bytes(path.read_bytes(), policy=policy.default)
            body = parsed.get_body(preferencelist=("plain", "html"))
            parts = [
                f"From: {parsed.get('From', '')}",
                f"To: {parsed.get('To', '')}",
                f"Subject: {parsed.get('Subject', '')}",
                f"Date: {parsed.get('Date', '')}",
                body.get_content() if body else "",
            ]
        except Exception as exc:  # noqa: BLE001
            return _unavailable(item_id, Backend.EMAIL, f"unreadable:{type(exc).__name__}")

    text = _clip("\n".join(str(part) for part in parts), options)
    if not text.strip():
        return _unavailable(item_id, Backend.EMAIL, "message_has_no_text")
    return Extraction(item_id=item_id, backend=str(Backend.EMAIL), text=text, char_count=len(text))


_BACKENDS = {
    Backend.PDF_TEXT: _pdf_text,
    Backend.PDF_OCR: _pdf_ocr,
    Backend.DOCX: _docx,
    Backend.XLSX: _xlsx,
    Backend.EMAIL: _email,
    Backend.IMAGE_OCR: _image_ocr,
}


def chain_for(backend: Backend, media_type: str) -> tuple[Backend, ...]:
    """Every backend that could run for this item, in the order they would be tried.

    Empty for a media type `auto_chain` does not cover.  Resolving the chain up front is
    what lets `extract_many` decide, without opening the file, whether an item is going to
    reach lxml.
    """
    if backend is not Backend.AUTO_CHAIN:
        return (backend,)
    return _CHAIN.get(media_type, ())


def extract(
    *,
    item_id: str,
    path: Path,
    media_type: str,
    backend: Backend = Backend.AUTO_CHAIN,
    options: ExtractOptions | None = None,
) -> Extraction:
    options = options or ExtractOptions()

    if backend is not Backend.AUTO_CHAIN:
        return _BACKENDS[backend](item_id, path, options)

    chain = chain_for(backend, media_type)
    if not chain:
        return _unavailable(item_id, Backend.AUTO_CHAIN, f"unsupported_media_type:{media_type}")

    last = _unavailable(item_id, Backend.AUTO_CHAIN, "no_backend_attempted")
    for candidate in chain:
        last = _BACKENDS[candidate](item_id, path, options)
        if last.usable:
            return last
    return last


# --------------------------------------------------------------------------------------
# Extracting a corpus
# --------------------------------------------------------------------------------------


#: Backends that reach lxml — python-docx and openpyxl directly, extract-msg through
#: BeautifulSoup.  ADR 0006 accepted running lxml under a pinned-off GIL *because there
#: were no threads*; its thread-safety is undeclared, and lxml 7 is still in beta.  Now
#: that threads exist, these keep running one at a time, and the deal ADR 0006 struck
#: holds. Delete this set the day lxml declares free-threaded support.
_LXML_BACKENDS = frozenset({Backend.DOCX, Backend.XLSX, Backend.EMAIL})


def reaches_lxml(backend: Backend, media_type: str) -> bool:
    """Whether extracting this item could reach lxml, counting fallbacks not yet taken.

    Conservative on purpose: a PDF that falls back to OCR never touches lxml, but a chain
    containing one lxml backend is treated as lxml-bound whether or not it gets that far.
    """
    return any(candidate in _LXML_BACKENDS for candidate in chain_for(backend, media_type))


#: Measured, not chosen — and measured because one-per-CPU was *slower*.  Over 120 padded
#: invoice PDFs on a 15-CPU host: 1 worker 3.56s, 2 workers 1.50x, 4 workers 2.29x, then
#: it regresses — 8 workers 1.98x, 15 workers 1.76x.  The interpreter is not the ceiling;
#: the same host scales arithmetic 6.8x and allocation-heavy work 4.3x at 15 threads.
#: pdfminer is: `PSLiteralTable` and `PSKeywordTable` are module-level singletons whose
#: `intern` does a check-then-insert into one shared dict for every token in every
#: document, so past a handful of threads the parsers spend their time queueing on it.
#:
#: This is the shape ADR 0005 calls a measured constant, and it belongs in the measurement
#: store rather than here — the knee moves with the host and with the corpus.  It is a
#: literal today because there is one measurement, taken on one machine.
_WORKER_CEILING = 4


def worker_count() -> int:
    """How many extractions to run at once: one per CPU, capped at the measured knee.

    Not a configuration slot.  A domain subagent selects an `extract=` group from the
    closed catalog; how many threads the host can spare is a property of the machine, not
    a decision a model gets to make, and nothing in a plan should differ between two hosts
    that ran the same composition.
    """
    return max(1, min(os.process_cpu_count() or 1, _WORKER_CEILING))


def _extract_one(item: Item, backend: Backend, options: ExtractOptions) -> Extraction:
    """`extract`, with the guarantee that it returns rather than raises.

    Each backend already turns a corrupt document into a reason, but a worker that raised
    anyway would take an item out of the batch entirely, and accounting has to balance:
    every item leaves with text or with a reason.
    """
    try:
        return extract(
            item_id=item.item_id,
            path=item.path,
            media_type=item.media_type,
            backend=backend,
            options=options,
        )
    except Exception as exc:  # noqa: BLE001 - an item that raised is still an item
        return _unavailable(item.item_id, backend, f"unreadable:{type(exc).__name__}")


def extract_many(
    items: Sequence[Item],
    *,
    backend: Backend = Backend.AUTO_CHAIN,
    options: ExtractOptions | None = None,
    max_workers: int | None = None,
) -> tuple[Extraction, ...]:
    """Extract a corpus, in parallel where that is safe, in input order always.

    Extraction is share-nothing per item — open a file, read it, return text — so it is
    the work ADR 0006 named as the first thing to parallelise, and at corpus scale it is
    where nearly all of a run's serial time goes.

    Two things are deliberately not left to the scheduler:

    - **Order.** Results are placed by input index, not by completion, because a plan is
      portable data: two runs over the same tree must produce the same inventory in the
      same order, and thread completion order is neither stable nor reproducible.
    - **lxml.** Items whose backend chain could reach lxml run first, serially, on this
      thread, before any worker exists.  See `_LXML_BACKENDS`.
    """
    options = options or ExtractOptions()
    items = tuple(items)
    results: list[Extraction | None] = [None] * len(items)

    lxml_bound: list[int] = []
    threadable: list[int] = []
    for index, item in enumerate(items):
        target = lxml_bound if reaches_lxml(backend, item.media_type) else threadable
        target.append(index)

    # First, and alone: no worker thread exists while these run.
    for index in lxml_bound:
        results[index] = _extract_one(items[index], backend, options)

    workers = worker_count() if max_workers is None else max(1, max_workers)
    if workers == 1 or len(threadable) <= 1:
        for index in threadable:
            results[index] = _extract_one(items[index], backend, options)
    else:
        # One document on this thread before any worker exists. pdfminer interns every
        # name it parses into two module-level tables with a check-then-insert, so two
        # threads meeting an un-interned name can each build a distinct object for it.
        # Parsing one document first interns the standard PDF names while nothing is
        # racing; what remains is document-specific and compared by value, not identity.
        first, rest = threadable[0], threadable[1:]
        results[first] = _extract_one(items[first], backend, options)
        with ThreadPoolExecutor(
            max_workers=min(workers, len(rest)), thread_name_prefix="doc-extract"
        ) as pool:
            pending = {
                pool.submit(_extract_one, items[index], backend, options): index
                for index in rest
            }
            for future in as_completed(pending):
                results[pending[future]] = future.result()

    missing = [index for index, result in enumerate(results) if result is None]
    if missing:  # pragma: no cover - every index is in exactly one of the two lists
        raise RuntimeError(f"extraction lost items at indices {missing}")
    return tuple(result for result in results if result is not None)
